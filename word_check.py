#!/usr/bin/env python
# coding: utf-8

import docx
import xlsxwriter
import xml.etree.ElementTree as ET
import os
import shutil
import re
import xlrd
import csv
from copy import deepcopy
import ast
import numpy as np
import math
import itertools

def clean_string(s):
    # 移除全形空格
    s = s.replace('\u3000', '')
    
    # 移除 (cm)
    s = s.replace('(cm)', '')
    
    # 移除換行符號
    s = s.replace('\n', '')
    
    # 移除多餘的空格
    s = s.strip()
    
    return s

def clear_type(type_list):
    new_type_list = []
    for t in type_list:
        t = t.split('-')[0].replace('TYPE', 'Type').replace('A', '')
        if t not in new_type_list:
            new_type_list.append(t)

    return len(new_type_list), new_type_list

def read_word(wordName, designFile, type_list):
    print('抓取設計計算書')
    doc = docx.Document(wordName)
    num_workItemType_design = 0
    count = 0
    count_blank = 0
    sheet_list = []

    count, sheet_list = clear_type(type_list)
    num_workItemType_design = count

    for i in range(num_workItemType_design - 1):
        designFile.append( deepcopy(designFile[0]) )
        i += 1

    for i in range(num_workItemType_design):
        element = designFile[i]
        element.set('TYPE', sheet_list[i])

    #找混凝土強度
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("混凝土強度")!=-1:
            arr.append(re.sub('[^0-9]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    for i in range(len(arr)):
        if i%(len(arr)/num_workItemType_design)==0:
            #print(arr[i])
            value = designFile[int(i/(len(arr)/num_workItemType_design))].find('Concrete/Strength/Value')
            #print('設計', value.get('unit'), value.text)
            value.text = arr[i]

    #找擋土壁深度
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("結論：連續壁深度採用")!=-1:
            arr.append(re.sub('[^0-9]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    for i in range(len(arr)):
        if i%(len(arr)/num_workItemType_design)==0:
            #print(arr[i])
            value = designFile[int(i/(len(arr)/num_workItemType_design))].find('Concrete/Depth/Value')
            #print('設計', value.get('unit'), value.text)
            value.text = arr[i]

    #找擋土壁厚度
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("擋土壁厚度")!=-1:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    for i in range(len(arr)):
        if i%(len(arr)/num_workItemType_design)==0:
            #print(arr[i])
            value = designFile[int(i/(len(arr)/num_workItemType_design))].find('Concrete/Thickness/Value')
            #print('設計', value.get('unit'), value.text)
            value.text = arr[i]

    #找中間柱長度
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("中間柱長度")!=-1:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    for i in range(len(arr)):
        if i%(len(arr)/(num_workItemType_design + count_blank))==0:
            #print(arr[i])
            value = designFile[int(i/(len(arr)/num_workItemType_design))].find('MiddleColumn/Length/Value')
            value.text = arr[i]
    #         print('設計', value.get('unit'), value.text)

    #找開挖深度
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("開挖深度")!=-1:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    for i in range(len(arr)):
        if i%(len(arr)/(num_workItemType_design + count_blank))==1: #這邊特別用了1
    #         print(arr)
            value = designFile[int(i/(len(arr)/num_workItemType_design))].find('MiddleColumn/Depth/Value')
            value.text = arr[i]
    #         print('設計', value.get('unit'), value.text)

    #找型鋼尺寸
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("型鋼尺寸")!=-1:
            textP = paragraph.text
            arr.append(textP[textP.find('=')+2:])
    #         print(re.sub('[^0-9]', '', paragraph.text))
    for i in range(len(arr)):
        if i%(len(arr)/(num_workItemType_design + count_blank))==0:
    #         print(arr[i])
            value = designFile[int(i/(len(arr)/num_workItemType_design))].find('MiddleColumn/Steel/Type/Value')
            value.text = arr[i]
    #         print('設計', value.text)

    #找型鋼長度
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("型鋼長度")!=-1:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    for i in range(len(arr)):
        if i%(len(arr)/(num_workItemType_design + count_blank))==0:
            #print(arr[i])
            value = designFile[int(i/(len(arr)/num_workItemType_design))].find('MiddleColumn/Steel/Length/Value')
            value.text = arr[i]
    #         print('設計', value.get('unit'), value.text)

    #找鑽掘樁直徑
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("鑽掘樁直徑")!=-1:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    for i in range(len(arr)):
        if i%(len(arr)/(num_workItemType_design + count_blank))==0:
            #print(arr[i])
            value = designFile[int(i/(len(arr)/num_workItemType_design))].find('MiddleColumn/DrilledPile/Diameter/Value')
            value.text = arr[i]
    #         print('設計', value.get('unit'), value.text)

    #找樁身埋入深度
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("樁身埋入深度")!=-1:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    for i in range(len(arr)):
        if i%(len(arr)/(num_workItemType_design + count_blank))==0:
            #print(arr[i])
            value = designFile[int(i/(len(arr)/num_workItemType_design))].find('MiddleColumn/DrilledPile/Length/Value')
            value.text = arr[i]
    #         print('設計', value.get('unit'), value.text)

    def writeSupFen(support,no_arr,num_arr,type_arr):
        for i in range(len(no_arr)-2):
            support.append( deepcopy(support[0]) )
            i+=1
        for i in range(len(no_arr)-1):
            support[i].find('Layer/Value').text = no_arr[i+1]
            support[i].find('Type/Value').text = type_arr[i+1]
            support[i].find('Count/Value').text = num_arr[i+1]
            i+=1

    #找支撐
    #支撐階數

    count = 0
    for table in doc.tables:
        no = []
        num_1 = []
        type_1 = []
        num_2 = []
        type_2 = []

        try:
            if table.rows[0].cells[2].text=="支撐型號" and table.rows[0].cells[5].text=="圍令型號":
                for cell in table.columns[0].cells:
                    if cell.text == '(續)':
                        break;
                    no.append(cell.text)
                for cell in table.columns[1].cells:
                    if cell.text == '(續)':
                        break;                
                    num_1.append(cell.text)
                for cell in table.columns[2].cells:
                    if cell.text == '(續)':
                        break;                
                    type_1.append(cell.text)
                for cell in table.columns[4].cells:
                    if cell.text == '(續)':
                        break;                
                    num_2.append(cell.text)
                for cell in table.columns[5].cells:
                    if cell.text == '(續)':
                        break;                
                    type_2.append(cell.text)                
    #             print(no)
    #             print(num_1)
    #             print(type_1)
    #             print(num_2)
    #             print(type_2)
                writeSupFen(designFile[count][1],no,num_1,type_1)
                writeSupFen(designFile[count][2],no,num_2,type_2)
                count += 1
        except:
            pass

    def writeVertical(node,start_list,end_list,type_list):
        for i in range(len(start_list)-2):
            node.append( deepcopy(node[0]) )
            i+=1
        for i in range(len(start_list)-1):
            node[i].find('DepthStart/Value').text = start_list[i+1]
            node[i].find('DepthEnd/Value').text = end_list[i+1]
            node[i].find('Type/Value').text = type_list[i+1]
            i+=1

    count = 0
    for table in doc.tables:
        start_list = []
        end_list = []
        type_list = []

        try:
            if table.rows[0].cells[3].text=="垂直筋設計":
                for cell in table.columns[0].cells:
                    start_list.append(cell.text)
                for cell in table.columns[1].cells:              
                    end_list.append(cell.text)
                for cell in table.columns[3].cells:             
                    type_list.append(clean_string(cell.text))
                writeVertical(designFile[math.floor(count/2)][3][0][count%2],start_list,end_list,type_list)
    #             print(count)
    #             print(count%2)
    #             print(math.floor(count/2))
                # print(designFile[math.floor(count/2)][3][0][count%2])
                # print(start_list)
                # print(end_list)
                # print(type_list)                        
                count += 1
        except:
            pass

    count = 0
    for table in doc.tables:
        start_list = []
        end_list = []
        type_list = []

        try:
            if table.rows[0].cells[3].text=="剪力筋設計":
                for cell in table.columns[0].cells:
                    start_list.append(cell.text)
                for cell in table.columns[1].cells:              
                    end_list.append(cell.text)
                for cell in table.columns[3].cells:             
                    type_list.append(clean_string(cell.text))
                writeVertical(designFile[count][3][2],start_list,end_list,type_list)
                # print(count)
                # print(designFile[count][3][2])
                # print(start_list)
                # print(end_list)
                # print(type_list)                        
                count += 1
        except:
            pass

    #水平筋
    arr = []
    parseStart = False
    for paragraph in doc.paragraphs:
        if paragraph.text.find("水平筋")!=-1:
            parseStart = True
        elif paragraph.text.find("剪力筋")!=-1:
            parseStart = False
        elif parseStart:
            if paragraph.text.find("@")!=-1:
                arr.append(re.sub('[^0-9@D]', '', paragraph.text.split('(cm)')[0]))

    for i in range(len(arr)):
        if i%(len(arr)/num_workItemType_design)==0:
            #print(arr[i])
            value = designFile[int(i/(len(arr)/num_workItemType_design))].find('RebarGroup/HorznRebar/Rebar/Type/Value')
            #print('設計', value.get('unit'), value.text)
            value.text = arr[i]

    print('設計計算書抓取完成')
    return sheet_list, num_workItemType_design
