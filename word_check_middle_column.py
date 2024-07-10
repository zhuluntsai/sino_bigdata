#!/usr/bin/env python
# coding: utf-8

import docx
import xlsxwriter
import xml.etree.ElementTree as ET
import os
import shutil
import re
import xlrd
import csv
from copy import deepcopy
import ast
import numpy as np
import math
import itertools

#清理字串
def clean_string(s):
    # 移除全形空格為空字符  #\u3000是全形空格在Unicode中的代碼
    s = s.replace('\u3000', '')
    
    # 移除 (cm) 為空字符
    s = s.replace('(cm)', '')
    
    # 移除換行符號為空字符 #將文本轉為單行
    s = s.replace('\n', '')
    
    # 移除多餘的空格(字符串兩端的空格) 
    s = s.strip()
    
    return s

#轉換字符與計數
def clear_type(count_blank, type_list):
    new_type_list = []
    ##type_list=['LG10連續壁100cm(TYPE T1)', 'LG10連續壁80cm(TYPE T2)', 'LG10排樁80cm(TYPE T3)']
    for t in type_list:
        if "空打" not in t:  # 剔除包含 "空打" 字樣的項目
            if not any(n in t for n in new_type_list):
                new_type_list.append(t)
            else:
                count_blank += 1
        #print('t:',t)
        #t = t.split('-')[0].replace('TYPE', 'Type')  
        #print('TYPE轉換:',t)
        ##if not any(n in t for n in new_type_list): 
        ##    new_type_list.append(t)
        ##else: 
        ##    count_blank += 1 
    print('new_type_list為:',new_type_list)
    return len(new_type_list), count_blank, new_type_list # new_type_list 表處理後不重複的項目串列 ; count_blank 表重複項目的數量 

def read_word(wordName, designFile, type_list): 
    print('抓取設計計算書')
    doc = docx.Document(wordName)  
    num_workItemType_design = 0 
    count = 0 
    count_blank = 0 
    sheet_list = [] #sheet_list = new_type_list，就是不包含重複項目的串列。

    count, count_blank, sheet_list = clear_type(count_blank, type_list)  #sheet_list [['80cmφ擋土排樁']]-->['LG10連續壁100cm(TYPE T1)', 'LG10連續壁80cm(TYPE T2)', 'LG10排樁80cm(TYPE T3)']
    #count = len(new_type_list) 不重複項目的數量  #count_blank = count_blank 重複項目的數量   #sheet_list = new_type_list  不重複項目的串列
    num_workItemType_design = count  #num_workItemType_design 不重複項目的數量
    print('count_blank為:',count_blank) 
    print('num_workItemType_design為:',num_workItemType_design) 
    print('sheet_list為:',sheet_list)

    #deepcopy出其他的連續壁
    for i in range(num_workItemType_design -1 ): 
        designFile.append( deepcopy(designFile[0]) ) 
        i += 1
        #print(i)

    for i in range(num_workItemType_design): 
        element = designFile[i] 
        element.set('TYPE', sheet_list[i]) 
        #print(i)
    
    # 排樁型式(LG10)
    arr = []
    for paragraph in doc.paragraphs:
        if "Type" in paragraph.text and "排樁" in paragraph.text:
            matches = re.findall(r'Type\s+\w+\s*排樁', paragraph.text) 
            for match in matches:
                sheet_pile_type = re.search(r'Type\s+(\w+)\s*排樁', match).group().replace('排樁', '') # 提取 "Type" 後面的字元
                arr.append(sheet_pile_type)
    print('排樁型式:', arr)
    
    for sheet_pile_type in arr: 
        matched_type = None
        for type_value in sheet_list:
            # 比對 TYPE 和 TYPE 後面的字母及數字
            type_suffix = type_value.split('(TYPE ')[1].split(')')[0].replace(" ", "")
            sheet_suffix = sheet_pile_type.split('Type')[1].replace(" ", "")
            if sheet_suffix in type_suffix:
                matched_type = type_value
                break
        if matched_type:
            value = designFile.find(f".//WorkItemType[@TYPE='{matched_type}']/Rowpile/Type/Value")
            if value is not None:
                value.text = sheet_pile_type
                print(f"Set value for TYPE '{matched_type}': {value.text}")
            else:
                print(f"No matching node found for TYPE '{matched_type}'")
        else:
            print(f"No matching TYPE found for sheet pile type '{sheet_pile_type}'")
    
    #排樁高度(LG10)
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("結論：排樁深度採用")!=-1:
            arr.append(re.sub('[^0-9.]', '', paragraph.text)) 
            #print(re.sub('[^0-9]', '', paragraph.text))
    print("排樁高度",arr)

    for height in arr:
        matched_type = None
        for type_value in sheet_list:
            # 比對節點名稱和關鍵字中的字樣
            if "排樁" in type_value:
                matched_type = type_value
                break
        if matched_type:
            value = designFile.find(f".//WorkItemType[@TYPE='{matched_type}']/Rowpile/Height/Value")
            if value is not None:
                value.text = height
                print(f"Set value for TYPE '{matched_type}': {value.text}")
            else:
                print(f"No matching node found for TYPE '{matched_type}'")
        else:
            print(f"No matching TYPE found for height '{height}'")
    
    #排樁樁徑(LG10)
    arr = []
    for paragraph in doc.paragraphs:
     if "排樁樁徑：" in paragraph.text:
        match = re.search(r'：(\d+)', paragraph.text)
        if match:
            arr.append(match.group(1))
    print("排樁樁徑",arr)
    for diameter in arr:
        matched_type = None
        for type_value in sheet_list:
            # 比對節點名稱和關鍵字中的字樣
            if "排樁" in type_value:
                matched_type = type_value
                break
        if matched_type:
            value = designFile.find(f".//WorkItemType[@TYPE='{matched_type}']/Rowpile/Diameter/Value")
            if value is not None:
                value.text = diameter
                print(f"Set value for TYPE '{matched_type}': {value.text}")
            else:
                print(f"No matching node found for TYPE '{matched_type}'")
        else:
            print(f"No matching TYPE found for diameter '{diameter}'")
    

    #----------------------------

    # 連續壁
    # 混凝土強度
    strength_nodes = designFile.findall(".//DiaphragmWall/Concrete/Strength/Value")
    print(f"Found {len(strength_nodes)} 連續壁 strength nodes")

    arr = []
    for paragraph in doc.paragraphs:
        if "混凝土強度" in paragraph.text:
            arr.append(re.sub('[^0-9]', '', paragraph.text))
            # print(re.sub('[^0-9]', '', paragraph.text))
    print("混凝土強度＿連續壁", arr)

    # 過濾出每個索引為 0, 2, 4, 6 ... 的數據
    filtered_arr = [arr[i] for i in range(0, len(arr), 2)]

    # 寫入到對應的 XML 節點
    count = 0
    for type_node in designFile.findall(".//WorkItemType"):
        type_value = type_node.get("TYPE")
        if "連續壁" in type_value:
            strength_node = type_node.find(".//DiaphragmWall/Concrete/Strength/Value")
            if strength_node is not None and count < len(filtered_arr):
                strength_node.text = filtered_arr[count]
                print(f"Set value for {type_value}: {strength_node.text}")
                count += 1
            else:
                print(f"No matching node found or no more values to assign for {type_value}")
    
    # 找擋土壁深度(LG10)
    arr = []
    for paragraph in doc.paragraphs:
        if "結論：連續壁深度採用" in paragraph.text:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            # print(re.sub('[^0-9]', '', paragraph.text))
    print("連續壁深度", arr)

    # 寫入到對應的 XML 節點
    count = 0
    for type_node in designFile.findall(".//WorkItemType"):
        type_value = type_node.get("TYPE")
        if "連續壁" in type_value:
            depth_node = type_node.find(".//DiaphragmWall/Depth/Value")
            if depth_node is not None and count < len(arr):
                depth_node.text = arr[count]
                print(f"Set depth value for {type_value}: {depth_node.text}")
                count += 1
            else:
                print(f"No matching node found or no more values to assign for {type_value}")
    
    # 找擋土壁厚度(LG10)
    arr = []
    for paragraph in doc.paragraphs:
        if "擋土壁厚度" in paragraph.text:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    print("擋土壁厚度", arr)

    # 抓取第一、第四、第七、第十...數值
    selected_values = []
    for i in range(0, len(arr), 3):
        if i < len(arr):
            selected_values.append(arr[i])
    print("選擇的擋土壁厚度", selected_values)

    # 將數據寫入有 "連續壁" 字樣的 TYPE 中
    selected_index = 0
    for work_item_type in designFile.findall(".//WorkItemType"):
        if "連續壁" in work_item_type.get("TYPE"):
            thickness_node = work_item_type.find(".//DiaphragmWall/Thickness/Value")
            if thickness_node is not None and selected_index < len(selected_values):
                thickness_node.text = selected_values[selected_index]
                selected_index += 1
                print(f"Set value for {work_item_type.get('TYPE')}: {thickness_node.text}")

    """
    # 鋼筋強度（LG10)
    arr_greater = []
    arr_less = []

    # 抓取鋼筋強度數值
    for paragraph in doc.paragraphs:
        if "鋼筋強度" in paragraph.text:
            matches_greater = re.findall(r'[≧>=]\s*.*?fy\s*=\s*(\d+)', paragraph.text)
            matches_less = re.findall(r'[<]\s*.*?fy\s*=\s*(\d+)', paragraph.text)   
            for match in matches_greater:
                arr_greater.append(match)
            for match in matches_less:
                arr_less.append(match)

    print("連續壁＿鋼筋強度(≧)", arr_greater)
    print("連續壁＿鋼筋強度(<)", arr_less)

    def write_values_to_xml(arr, description):
        # 找到有 "連續壁" 字樣的節點
        wall_nodes = [node for node in designFile.findall(".//WorkItemType") if "連續壁" in node.get("TYPE")]
        
        index = 0
        for i in range(len(wall_nodes)):
            if index < len(arr):
                node = wall_nodes[i].find(f".//RebarWeight[@Description='{description}']/Value")
                if node is not None:
                    node.text = arr[index]
                    print(f"Set value for {wall_nodes[i].get('TYPE')} - {description}: {node.text}")
                    index += 3  # 依次取第1、4、7、10...個數據
                else:
                    print(f"No matching node found for {wall_nodes[i].get('TYPE')} - {description}")

    # 將數值寫入 XML
    write_values_to_xml(arr_greater, "SD420W鋼筋噸數")
    write_values_to_xml(arr_less, "SD280W鋼筋噸數")
    """
    #----------------------------
    
    #混凝土強度(Y39)＿排樁
    arr = []
    contains = False

    # 先檢查是否包含 "排樁" 關鍵字
    for paragraph in doc.paragraphs:
        if "排樁" in paragraph.text:
            contains = True
            break

    # 抓取混凝土強度數值
    for paragraph in doc.paragraphs:
        if "混凝土強度" in paragraph.text:
            arr.append(re.sub('[^0-9]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))

    print("混凝土數值", arr)
    for i in range(len(arr)):
        if i % (len(arr) / num_workItemType_design) == 0:  
            # 抓垂直筋的 2/1=2 餘數必須為零，故只抓取串列序列為0的值(第一個數值，為垂直筋)
            # 抓垂直筋的 4/4/3 
            print(arr[i])
            value = None
            if contains:
                # 若包含 "排樁" 字樣，將數值存入 "LG10排樁80cm(TYPE T3)" 節點中
                for design_node in designFile.findall(".//WorkItemType"):
                    if "排樁" in design_node.get("TYPE"):
                        value = design_node.find('Concrete/Strength/Value')
                        break
            else:
                # 若不包含 "排樁" 字樣，根據原來的邏輯存入對應的節點中
                value = designFile[int(i / (len(arr) / num_workItemType_design))].find('Concrete/Strength/Value')
            
            if value is not None:
                value.text = arr[i]
                print(f"Set value for {'排樁' if contains else 'other'}: {value.text}")
            else:
                print("No matching node found for the specified type")

    # 鋼筋強度(Y39)＿排樁
    arr_greater = []
    arr_less = []
    contains = False

    # 先檢查是否包含 "排樁" 關鍵字
    for paragraph in doc.paragraphs:
        if "排樁" in paragraph.text:
            contains = True
            break

    # 抓取鋼筋強度數值
    for paragraph in doc.paragraphs:
        if "鋼筋強度" in paragraph.text:
            matches_greater = re.findall(r'[≧>=]\s*.*?fy\s*=\s*(\d+)', paragraph.text)
            matches_less = re.findall(r'[<]\s*.*?fy\s*=\s*(\d+)', paragraph.text)   
            for match in matches_greater:
                arr_greater.append(match)
            for match in matches_less:
                arr_less.append(match)

    print("鋼筋強度(≧)", arr_greater)
    print("鋼筋強度(<)", arr_less)
    if arr_greater:
        value = None
        if contains:
            # 若包含 "排樁" 字樣，將數值存入 "LG10排樁80cm(TYPE T3)" 節點中
            for design_node in designFile.findall(".//WorkItemType"):
                if "排樁" in design_node.get("TYPE"):
                    value = design_node.find(".//Strength[@Description='大於等於']/Value")
                    break
        else:
            # 若不包含 "排樁" 字樣，根據原來的邏輯存入對應的節點中
            value = designFile[0].find(".//Strength[@Description='大於等於']/Value")
        
        if value is not None:
            value.text = arr_greater[0]  # 只取第一個項目
            print(f"Set value for {'排樁' if contains else 'other'}: {value.text}")
        else:
            print("No matching node found for the specified type")

    if arr_less:
        value = None
        if contains:
            # 若包含 "排樁" 字樣，將數值存入 "LG10排樁80cm(TYPE T3)" 節點中
            for design_node in designFile.findall(".//WorkItemType"):
                if "排樁" in design_node.get("TYPE"):
                    value = design_node.find(".//Strength[@Description='小於']/Value")
                    break
        else:
            # 若不包含 "排樁" 字樣，根據原來的邏輯存入對應的節點中
            value = designFile[0].find(".//Strength[@Description='小於']/Value")
        
        if value is not None:
            value.text = arr_less[0]  # 只取第一個項目
            print(f"Set value for {'排樁' if contains else 'other'}: {value.text}")
        else:
            print("No matching node found for the specified type")
    
    #支撐圍令(LG10)
    #支撐階數
    def writeSupFen(support,no_arr,num_arr,type_arr):  
        for i in range(len(no_arr)-2):
            support.append( deepcopy(support[0]) )
            i+=1
            #print("支撐copyi",i)   
        for i in range(len(no_arr)-1):
            #print("支撐寫入i",i)
            support[i].find('Layer/Value').text = no_arr[i+1]
            support[i].find('Type/Value').text = type_arr[i+1]
            support[i].find('Count/Value').text = num_arr[i+1]
            i+=1
    count = 0
    valid_tables = []
    for table in doc.tables:
        no = []
        num_1 = []
        type_1 = []
        num_2 = []
        type_2 = []
        try:
            if table.rows[0].cells[2].text=="支撐型號" and table.rows[0].cells[5].text=="圍令型號":
                for cell in table.columns[0].cells:
                    if cell.text == '(續)':
                        break;
                    no.append(cell.text)
                for cell in table.columns[1].cells:
                    if cell.text == '(續)':
                        break;                
                    num_1.append(cell.text)
                for cell in table.columns[2].cells:
                    if cell.text == '(續)':
                        break;                
                    type_1.append(cell.text)
                for cell in table.columns[4].cells:
                    if cell.text == '(續)':
                        break;                
                    num_2.append(cell.text)
                for cell in table.columns[5].cells:
                    if cell.text == '(續)':
                        break;                
                    type_2.append(cell.text)   
                
                writeSupFen(designFile[count][3],no,num_1,type_1)
                writeSupFen(designFile[count][4],no,num_2,type_2)
                count += 1

                #valid_tables.append((no, num_1, type_1, num_2, type_2)) 
                #count += 1   
                #print(no)
                #print(num_1)
                #print(type_1)
                #print(num_2)
                #print(type_2)                  
        except:
            pass
    #print(valid_tables) 
    ##count = 0               
    ##no, num_1, type_1, num_2, type_2 = valid_tables[2]  #只處理第三個表格
    #print(no)
    #print(num_1)
    #print(type_1)
    #print(num_2)
    #print(type_2)  
    #print(count)
    #print("支撐len",len(designFile[0][3]))
    #print("圍令len",len(designFile[0][4]))

    ##writeSupFen(designFile[0][3], no, num_1, type_1) 
    ##writeSupFen(designFile[0][4], no, num_2, type_2) 
    

    """
    #拷貝中間柱
    middle_column_list = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("中間柱設計檢核")!=-1:
            middle_column_list.append(paragraph.text)
            #print(re.sub('[^0-9]', '', paragraph.text))
    print('中間柱設計檢核:',middle_column_list)

    middleColumnGroup = designFile.find(f"./*[@TYPE='{type_list[0]}']").find('MiddleColumnGroup')
    for i in range(len(middle_column_list)-1):
        middleColumnGroup.insert(0, deepcopy(middleColumnGroup[0]))
        print("設計中間柱拷貝",i)
    """

    #中間柱長度(LG10)＿中間樁＿v1
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("中間柱長度")!=-1:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    #print('中間柱長度:',arr)
    middle_column_nodes = designFile.findall('.//MiddleColumnGroup/MiddleColumn/Length/Value')
    print(len(middle_column_nodes))
    
    if len(middle_column_nodes) >= len(arr):
        for i in range(len(arr)):
            value = middle_column_nodes[i]
            value.text = arr[i]
            print('設計', value.get('unit'), value.text)
    else:
        print("Error: Not enough MiddleColumn nodes to assign values.")

    
    #開挖深度(LG10)
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("開挖深度 =")!=-1:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    #print('開挖深度:',arr)
    middle_column_nodes = designFile.findall('.//MiddleColumnGroup/MiddleColumn/Depth/Value')
    
    if len(middle_column_nodes) >= len(arr):
        for i in range(len(arr)):
            value = middle_column_nodes[i]
            value.text = arr[i]
            print('設計', value.get('unit'), value.text)
    else:
        print("Error: Not enough MiddleColumn nodes to assign values.")
    
    #型鋼尺寸(LG10)
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("型鋼尺寸")!=-1:
            textP = paragraph.text
            arr.append(textP[textP.find('=')+2:])
           #print(re.sub('[^0-9]', '', paragraph.text))
    #print('型鋼尺寸:',arr)
    middle_column_nodes = designFile.findall('.//MiddleColumnGroup/MiddleColumn/Steel/Type/Value')
    if len(middle_column_nodes) >= len(arr):
        for i in range(len(arr)):
            value = middle_column_nodes[i]
            value.text = arr[i]
            #print('設計', value.text)
    else:
        print("Error: Not enough MiddleColumn nodes to assign values.")
    
    #型鋼長度(LG10)
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("型鋼長度")!=-1:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    #print('型鋼長度:',arr)
    middle_column_nodes = designFile.findall('.//MiddleColumnGroup/MiddleColumn/Steel/Length/Value')
    #arr = arr[2:]
    if len(middle_column_nodes) >= len(arr):
        for i in range(len(arr)):
            value = middle_column_nodes[i]
            value.text = arr[i]
            #print('設計', value.get('unit'), value.text)
    else:
        print("Error: Not enough MiddleColumn nodes to assign values.")
    
    #鑽掘樁直徑(LG10)
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("鑽掘樁直徑")!=-1:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    #print('鑽掘樁直徑:',arr)
    middle_column_nodes = designFile.findall('.//MiddleColumnGroup/MiddleColumn/DrilledPile/Diameter/Value')
    #arr = arr[2:]
    if len(middle_column_nodes) >= len(arr):
        for i in range(len(arr)):
            value = middle_column_nodes[i]
            value.text = arr[i]
            #print('設計', value.get('unit'), value.text)
    else:
        print("Error: Not enough MiddleColumn nodes to assign values.")
    
    #樁身埋入深度(LG10)
    arr = []
    for paragraph in doc.paragraphs:
        if paragraph.text.find("樁身埋入深度")!=-1:
            arr.append(re.sub('[^0-9.]', '', paragraph.text))
            #print(re.sub('[^0-9]', '', paragraph.text))
    #print('樁身埋入深度:',arr)
    middle_column_nodes = designFile.findall('.//MiddleColumnGroup/MiddleColumn/DrilledPile/Length/Value')
    #arr = arr[2:]
    if len(middle_column_nodes) >= len(arr):
        for i in range(len(arr)):
            value = middle_column_nodes[i]
            value.text = arr[i]
            #print('設計', value.get('unit'), value.text)
    else:
        print("Error: Not enough MiddleColumn nodes to assign values.")

    #---------------------------------------    

    """
    #配筋設計(Y39)
    def writeVertical(node,start_list,end_list,type_list):
        for i in range(len(start_list)-2):
            node.append( deepcopy(node[0]) )
            i+=1
            #print("鋼筋copyi",i)
        for i in range(len(start_list)-1):
            #print("鋼筋寫入i",i)
            node[i].find('DepthStart/Value').text = start_list[i+1]
            node[i].find('DepthEnd/Value').text = end_list[i+1]
            node[i].find('Type/Value').text = type_list[i+1]
            i+=1  
    
    #垂直筋(Y39)
    count = 0
    for table in doc.tables:
        start_list = [] #開始深度
        end_list = [] #結束深度
        type_list = [] #垂直筋設計

        try:
            if table.rows[0].cells[3].text=="垂直筋設計":
                for cell in table.columns[0].cells:
                    start_list.append(cell.text)
                for cell in table.columns[1].cells:              
                    end_list.append(cell.text)
                for cell in table.columns[3].cells:             
                    type_list.append(clean_string(cell.text)) #clean_string
                writeVertical(designFile[math.floor(count/2)][1][2][count%2],start_list,end_list,type_list)
                #print(count)
                #print(count%2)
                #print(math.floor(count/2))
                #print(designFile[math.floor(count/2)][3][0][count%2])
                #print(start_list)
                #print(end_list)
                #print(type_list)  
                #print(count)                      
                count += 1
        except:
            pass   
    """
    """
    #剪力筋(Y39)
    count = 0
    for table in doc.tables:
        start_list = []
        end_list = []
        type_list = []

        try:
            if table.rows[0].cells[3].text=="剪力筋設計":
                for cell in table.columns[0].cells:
                    start_list.append(cell.text)
                for cell in table.columns[1].cells:              
                    end_list.append(cell.text)
                for cell in table.columns[3].cells:             
                    type_list.append(clean_string(cell.text))
                writeVertical(designFile[count][1][3],start_list,end_list,type_list) #schema結構
                # print(count)
                # print(designFile[count][1][3])
                #print(start_list)
                #print(end_list)
                #print(type_list)   
                #print(count)                     
                count += 1
        except:
            pass
    """
    print('設計計算書抓取完成')
    return sheet_list, num_workItemType_design
